#!/usr/bin/env python3
"""
DOCX Report Generation Module

This module generates DOCX reports from parsed Nessus data using finding tables
that match the exact format from finding.docx.
"""

import os
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DocxGenerator:
    """Generator for DOCX reports from Nessus vulnerability data"""
    
    def __init__(self, template_path: str = "report_template.docx"):
        self.template_path = template_path
    
    def generate_report(self, parsed_data: List[Dict[str, Any]], 
                       summary_stats: Dict[str, Any],
                       output_path: str,
                       include_informational: bool = True,
                       customer_abbreviation: str = "XXXXX",
                       network_type: str = "external") -> str:
        """
        Generate DOCX report with finding tables only (matching finding.docx format)
        
        Args:
            parsed_data: List of host dictionaries from NessusParser
            summary_stats: Summary statistics from NessusParser
            output_path: Path where to save the generated report
            include_informational: Whether to include informational level vulnerabilities
            customer_abbreviation: Customer abbreviation for PT codes
            network_type: Network type ('external' or 'internal') for PT codes
            
        Returns:
            Path to the generated DOCX file
        """
        try:
            # Create new document (no template needed for findings-only format)
            doc = Document()
            logging.info("Creating findings-only document")
            
            # Add only vulnerability finding tables
            self._add_finding_tables_only(doc, parsed_data, include_informational, customer_abbreviation, network_type)
            
            # Save the document
            doc.save(output_path)
            logging.info(f"Findings report generated: {output_path}")
            
            return output_path
            
        except Exception as e:
            logging.error(f"Error generating findings report: {e}")
            raise
    
    def _add_finding_tables_only(self, doc: Document, parsed_data: List[Dict[str, Any]], include_informational: bool = True, customer_abbreviation: str = "XXXXX", network_type: str = "external"):
        """Add finding tables in exact format matching finding.docx"""
        
        # Group vulnerabilities by plugin_id and plugin_name
        vulnerability_groups = {}
        
        # Debug logging
        import logging
        logging.info(f"Starting finding tables generation with {len(parsed_data)} hosts")
        
        for host_data in parsed_data:
            host_ip = host_data.get('host_ip', 'Unknown')
            host_fqdn = host_data.get('host_fqdn', '')
            host_os = host_data.get('os', '')
            
            for vuln in host_data.get('vulnerabilities', []):
                # Skip informational vulnerabilities if not requested
                if not include_informational and vuln.get('severity', '').lower() == 'informational':
                    continue
                
                plugin_id = vuln.get('plugin_id', '')
                plugin_name = vuln.get('plugin_name', '')
                vuln_key = f"{plugin_id}_{plugin_name}"
                
                # Debug CVE data type
                cve_data = vuln.get('cves', [])
                logging.debug(f"CVE data for {plugin_name}: {cve_data} (type: {type(cve_data)})")
                
                if vuln_key not in vulnerability_groups:
                    vulnerability_groups[vuln_key] = {
                        'plugin_id': plugin_id,
                        'plugin_name': plugin_name,
                        'severity': vuln.get('severity', ''),
                        'description': vuln.get('description', ''),
                        'solution': vuln.get('solution', ''),
                        'synopsis': vuln.get('synopsis', ''),
                        'cvss_score': vuln.get('cvss_score', ''),
                        'cvss_vector': vuln.get('cvss_vector', ''),
                        'cves': vuln.get('cves', []),
                        'risk_factor': vuln.get('risk_factor', ''),
                        'see_also': vuln.get('see_also', ''),
                        'plugin_output': vuln.get('plugin_output', ''),
                        'affected_hosts': []
                    }
                
                # Add affected host information
                host_info = {
                    'ip': host_ip,
                    'fqdn': host_fqdn,
                    'os': host_os,
                    'port': vuln.get('port', ''),
                    'protocol': vuln.get('protocol', '')
                }
                vulnerability_groups[vuln_key]['affected_hosts'].append(host_info)
        
        # Sort vulnerabilities by severity
        severity_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3, 'Informational': 4}
        sorted_vulns = sorted(
            vulnerability_groups.values(),
            key=lambda x: severity_order.get(x['severity'], 5)
        )
        
        # Log filtering results
        if not include_informational:
            informational_count = sum(1 for v in sorted_vulns if v['severity'].lower() == 'informational')
            logging.info(f"Filtered out {informational_count} informational vulnerabilities")
        
        logging.info(f"Generating {len(sorted_vulns)} finding tables")
        
        # Generate PT codes based on network type and customer abbreviation
        from datetime import datetime
        current_year = datetime.now().year
        current_month = datetime.now().month
        
        # Determine prefix based on network type
        code_prefix = "D" if network_type == "external" else "I"
        
        logging.info(f"Generating PT codes with format: PT-{customer_abbreviation}-{current_year}-{current_month:02d}-{code_prefix}XXX")
        
        # Create finding tables in exact format
        for i, vuln in enumerate(sorted_vulns, 1):
            # Add page break before each table (except the first one)
            if i > 1:
                doc.add_page_break()
            
            # Add heading for each table
            heading = doc.add_heading(level=2)
            heading.clear()  # Clear default content
            heading_run = heading.add_run(f"{i}. {vuln['plugin_name']}")
            heading_run.font.name = 'Calibri'
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            
            # Create table with 12 rows, 2 columns (matching finding.docx structure)
            table = doc.add_table(rows=12, cols=2)
            table.style = 'Normal Table'
            
            # Apply exact table properties from finding.docx
            self._apply_table_properties(table)
            
            # Set exact column widths (updated first column to 1.9 inches)
            table.columns[0].width = Inches(1.9)   # First column width set to 1.9 inches
            table.columns[1].width = Inches(4.91)  # Second column adjusted to maintain total width
            
            # Map severity to Turkish
            severity_map = {
                'Critical': 'KRİTİK SEVİYE',
                'High': 'YÜKSEK SEVİYE', 
                'Medium': 'ORTA SEVİYE',
                'Low': 'DÜŞÜK SEVİYE',
                'Informational': 'BİLGİLENDİRME'
            }
            
            # Prepare affected hosts list
            affected_hosts_text = ""
            for host in vuln['affected_hosts']:
                host_line = f"{str(host['ip'])}"
                if host['fqdn']:
                    host_line += f" ({str(host['fqdn'])})"
                if host['port']:
                    host_line += f":{str(host['port'])}/{str(host['protocol'])}"
                affected_hosts_text += host_line + "\n"
            
            # Generate PT code for this finding
            pt_code = f"PT-{customer_abbreviation}-{current_year}-{current_month:02d}-{code_prefix}{i:03d}"
            
            # Configure each row with exact formatting
            # Debug CVE handling before using
            cve_data = vuln['cves']
            logging.debug(f"Processing CVE for table: {cve_data} (type: {type(cve_data)})")
            
            # Safe CVE handling with comprehensive type checking
            try:
                if isinstance(cve_data, list) and cve_data:
                    cve_text = ', '.join(str(cve) for cve in cve_data)
                elif isinstance(cve_data, str) and cve_data:
                    cve_text = cve_data
                else:
                    cve_text = 'CVE bilgisi mevcut değil'
                logging.debug(f"Final CVE text: {cve_text}")
            except Exception as e:
                logging.error(f"Error processing CVE data {cve_data}: {e}")
                cve_text = 'CVE bilgisi mevcut değil'
            
            rows_data = [
                (pt_code, severity_map.get(vuln['severity'], 'BİLİNMEYEN SEVİYE')),
                ("Bulgu Adı", vuln['plugin_name']),
                ("CVE Kodu", cve_text),
                ("CVSS Skoru", f"{vuln['cvss_score'] if vuln['cvss_score'] else 'Skor bilgisi mevcut değil'}{' | ' + str(vuln['cvss_vector']) if vuln['cvss_vector'] else ''}"),
                ("Bulgu Türü", "Güvenlik Zafiyeti"),
                ("Bulgu Detayı", "Yapılan güvenlik testleri sonucunda ilgili sistemlerde belirtilen güvenlik zafiyetlerine rastlanmıştır.\n\nŞekil : Güvenlik zafiyeti tespitine yönelik ekran görüntüsü\n\nNOT: Ekran görüntüleri örneklem olarak paylaşılmıştır. Diğer sistemleri görmek için etkilenen bileşenlere bakınız."),
                ("Erişim Noktası", "İnternet"),
                ("Kullanıcı Profili", "Anonim Kullanıcı"),
                ("Etkilenen Bileşenler", affected_hosts_text.strip()),
                ("Bulgu Açıklaması", vuln['description'] or 'Açıklama bilgisi mevcut değil.'),
                ("Çözüm Önerisi", vuln['solution'] or 'Çözüm önerisi bilgisi mevcut değil.'),
                ("İlgili Bağlantılar", vuln['see_also'] or 'İlgili bağlantı bilgisi mevcut değil.')
            ]
            
            # Apply formatting to each row
            for row_idx, (label, content) in enumerate(rows_data):
                row = table.rows[row_idx]
                
                # Set row height to control table compactness
                tr = row._tr
                trPr = tr.trPr
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    tr.insert(0, trPr)
                trHeight = OxmlElement('w:trHeight')
                
                # Set compact minimum row heights
                # Use 'atLeast' rule so content can expand as needed
                min_heights = ['300', '300', '280', '280', '300', '600', '280', '250', '250', '800', '280', '350']
                if row_idx < len(min_heights):
                    trHeight.set(qn('w:val'), min_heights[row_idx])
                else:
                    trHeight.set(qn('w:val'), '280')  # Default minimum height
                
                # Set hRule to 'atLeast' so content can expand
                trHeight.set(qn('w:hRule'), 'atLeast')
                trPr.append(trHeight)
                
                # Set cell content
                row.cells[0].text = label
                row.cells[1].text = content
                
                # Apply exact formatting from finding.docx
                severity_text = content if row_idx == 0 else ''  # Pass severity for row 1
                self._format_table_cell(row.cells[0], is_header=True, row_idx=row_idx)
                self._format_table_cell(row.cells[1], is_header=False, row_idx=row_idx, severity_text=severity_text)
            
            # Add spacing between tables
            doc.add_paragraph()
            doc.add_paragraph()
    
    def _apply_table_properties(self, table):
        """Apply exact table properties from finding.docx"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Set table width (9786 dxa from finding.docx)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '9786')
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        # Set table indent (-147 dxa from finding.docx)
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '-147')
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        # Set table borders (single, 4pt, black)
        tblBorders = OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Set table layout to fixed
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # Set table cell margins (0 left/right from finding.docx)
        tblCellMar = OxmlElement('w:tblCellMar')
        for margin_type in ['left', 'right']:
            margin = OxmlElement(f'w:{margin_type}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tblCellMar.append(margin)
        tblPr.append(tblCellMar)
        
        # Set table look properties
        tblLook = OxmlElement('w:tblLook')
        tblLook.set(qn('w:val'), '01E0')
        tblLook.set(qn('w:firstRow'), '1')
        tblLook.set(qn('w:lastRow'), '1')
        tblLook.set(qn('w:firstColumn'), '1')
        tblLook.set(qn('w:lastColumn'), '1')
        tblLook.set(qn('w:noHBand'), '0')
        tblLook.set(qn('w:noVBand'), '0')
        tblPr.append(tblLook)
    
    def _format_table_cell(self, cell, is_header: bool, row_idx: int = 0, severity_text: str = ''):
        """Apply exact formatting to table cells matching finding.docx"""
        
        # Set cell properties
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        
        # Set vertical alignment to center
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        
        # Set cell width
        tcW = OxmlElement('w:tcW')
        if is_header:
            tcW.set(qn('w:w'), '2736')  # First column width (1.9 inches = 2736 dxa)
            # Add background color based on row and cell
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            if row_idx == 0:  # Row 1 has special colors
                shd.set(qn('w:fill'), '002060')  # Dark blue for first cell
            else:
                shd.set(qn('w:fill'), 'F2F2F2')  # Gray for other header cells
            tcPr.append(shd)
        else:
            tcW.set(qn('w:w'), '7050')  # Second column width (adjusted to maintain total width)
            # Add background color for Row 1, second column (severity-based)
            if row_idx == 0:  # Row 1, second column - severity cell
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                # Use passed severity_text to determine color
                if 'KRİTİK' in severity_text.upper():
                    shd.set(qn('w:fill'), 'FF0000')  # Red for Critical
                elif 'YÜKSEK' in severity_text.upper():
                    shd.set(qn('w:fill'), 'FFA500')  # Orange for High
                elif 'ORTA' in severity_text.upper():
                    shd.set(qn('w:fill'), 'FFFF00')  # Yellow for Medium
                elif 'DÜŞÜK' in severity_text.upper():
                    shd.set(qn('w:fill'), '00FF00')  # Green for Low
                elif 'BİLGİLENDİRME' in severity_text.upper():
                    shd.set(qn('w:fill'), '87CEEB')  # Sky blue for Informational
                else:
                    shd.set(qn('w:fill'), 'FFFF00')  # Default yellow
                tcPr.append(shd)
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        
        # Set consistent cell margins for all cells
        tcMar = OxmlElement('w:tcMar')
        for margin_type in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_type}')
            margin.set(qn('w:w'), '113')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
        
        # Format paragraphs in cell
        for paragraph in cell.paragraphs:
            # Set compact line spacing to reduce row height
            pPr = paragraph._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '240')
            spacing.set(qn('w:lineRule'), 'atLeast')
            # Remove spacing before and after paragraphs
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            
            # Set consistent paragraph indentation for all cells
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '113')
            ind.set(qn('w:right'), '113')
            pPr.append(ind)
            
            # Set specific spacing and alignment based on row and content
            if is_header:  # First column - all left aligned
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif row_idx == 0:  # Row 1, second column (center alignment)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif row_idx == 5:  # Bulgu Detayı row - has special formatting
                self._format_bulgu_detayi_paragraph(paragraph, spacing, pPr)
                continue  # Skip default spacing
            elif row_idx in [9, 10, 11]:  # Description, Solution, Links rows
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Keep compact spacing for these rows too
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
            
            pPr.append(spacing)
            
            # Format runs (text formatting)
            for run in paragraph.runs:
                if is_header:
                    run.font.bold = True
                
                # Force all fonts to be Calibri (override any Cambria)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    def _format_bulgu_detayi_paragraph(self, paragraph, spacing, pPr):
        """Apply special formatting to Bulgu Detayı cell content matching finding.docx exactly"""
        # Clear existing content and rebuild with proper formatting
        text_content = paragraph.text
        cell = paragraph._parent
        paragraph.clear()
        
        # Set alignment to both (justified) for main content
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Set compact spacing to reduce row height
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        pPr.append(spacing)
        
        # Split content into parts
        parts = text_content.split('\n\n')
        
        if len(parts) >= 1:
            # First part - main text (justified)
            run = paragraph.add_run(parts[0])
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            
            # Add figure caption as separate paragraph with center alignment
            if len(parts) > 1:
                figure_text = parts[1].strip()
                
                # Add 3 empty lines before figure caption
                paragraph.add_run('\n\n\n')
                
                # Create new paragraph for figure caption
                figure_para = cell.add_paragraph()
                figure_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Set spacing for figure paragraph
                figure_pPr = figure_para._element.get_or_add_pPr()
                figure_spacing = OxmlElement('w:spacing')
                figure_spacing.set(qn('w:line'), '240')
                figure_spacing.set(qn('w:lineRule'), 'atLeast')
                figure_spacing.set(qn('w:before'), '0')
                figure_spacing.set(qn('w:after'), '0')
                figure_pPr.append(figure_spacing)
                
                # Add figure text
                figure_run = figure_para.add_run(figure_text)
                figure_run.font.name = 'Calibri'
                figure_run.font.size = Pt(11)
                figure_run.font.italic = True
                figure_run.font.color.rgb = RGBColor(0x36, 0x5F, 0x91)  # Blue color from analysis
                
            # Add NOTE section as separate paragraph
            if len(parts) > 2:
                note_text = parts[2]
                
                # Create new paragraph for NOTE
                note_para = cell.add_paragraph()
                note_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Set spacing for note paragraph
                note_pPr = note_para._element.get_or_add_pPr()
                note_spacing = OxmlElement('w:spacing')
                note_spacing.set(qn('w:line'), '240')
                note_spacing.set(qn('w:lineRule'), 'atLeast')
                note_spacing.set(qn('w:before'), '0')
                note_spacing.set(qn('w:after'), '0')
                note_pPr.append(note_spacing)
                
                if note_text.startswith('NOT:'):
                    note_run1 = note_para.add_run('NOT:')
                    note_run1.font.name = 'Calibri'
                    note_run1.font.size = Pt(11)
                    note_run1.font.bold = True
                    note_run2 = note_para.add_run(note_text[4:])
                    note_run2.font.name = 'Calibri'
                    note_run2.font.size = Pt(11)
                else:
                    note_run = note_para.add_run(note_text)
                    note_run.font.name = 'Calibri'
                    note_run.font.size = Pt(11)

    def _add_vulnerability_tables(self, doc: Document, parsed_data: List[Dict[str, Any]]):
        """Add dynamic vulnerability tables grouped by vulnerability type"""
        
        # Find placeholder for vulnerability table or add at the end
        vuln_table_placeholder = "${VulnTable}"
        placeholder_found = False
        
        # Look for placeholder in existing content
        for paragraph in doc.paragraphs:
            if vuln_table_placeholder in paragraph.text:
                # Replace placeholder with actual table
                paragraph.clear()
                paragraph.add_run("Detailed Vulnerability Findings")
                placeholder_found = True
                break
        
        if not placeholder_found:
            # Add section header
            heading = doc.add_heading('Detailed Vulnerability Findings', level=2)
        
        # Group vulnerabilities by plugin_id and plugin_name
        vulnerability_groups = {}
        
        for host_data in parsed_data:
            host_ip = host_data.get('host_ip', 'Unknown')
            host_fqdn = host_data.get('host_fqdn', '')
            host_os = host_data.get('os', '')
            
            for vuln in host_data.get('vulnerabilities', []):
                plugin_id = vuln.get('plugin_id', '')
                plugin_name = vuln.get('plugin_name', '')
                vuln_key = f"{plugin_id}_{plugin_name}"
                
                if vuln_key not in vulnerability_groups:
                    vulnerability_groups[vuln_key] = {
                        'plugin_id': plugin_id,
                        'plugin_name': plugin_name,
                        'severity': vuln.get('severity', ''),
                        'description': vuln.get('description', ''),
                        'solution': vuln.get('solution', ''),
                        'synopsis': vuln.get('synopsis', ''),
                        'cvss_score': vuln.get('cvss_score', ''),
                        'cvss_vector': vuln.get('cvss_vector', ''),
                        'cves': vuln.get('cves', []),
                        'risk_factor': vuln.get('risk_factor', ''),
                        'affected_hosts': []
                    }
                
                # Add affected host information
                host_info = {
                    'ip': host_ip,
                    'fqdn': host_fqdn,
                    'os': host_os,
                    'port': vuln.get('port', ''),
                    'protocol': vuln.get('protocol', '')
                }
                vulnerability_groups[vuln_key]['affected_hosts'].append(host_info)
        
        # Sort vulnerabilities by severity (Critical, High, Medium, Low, Informational)
        severity_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3, 'Informational': 4}
        sorted_vulns = sorted(
            vulnerability_groups.values(),
            key=lambda x: severity_order.get(x['severity'], 5)
        )
        
        # Create detailed vulnerability sections
        for i, vuln in enumerate(sorted_vulns, 1):
            # Vulnerability heading
            vuln_heading = doc.add_heading(f"{i}. {vuln['plugin_name']}", level=3)
            
            # Basic vulnerability information table
            info_table = doc.add_table(rows=8, cols=2)
            info_table.style = 'Table Grid'
            
            # Populate vulnerability information
            links_text = vuln.get('see_also', '') or 'See description for additional references'
            info_data = [
                ('Plugin ID', vuln['plugin_id']),
                ('Vulnerability Name', vuln['plugin_name']),
                ('Severity', vuln['severity']),
                ('CVSS Score', vuln['cvss_score']),
                ('CVSS Vector', vuln['cvss_vector']),
                ('CVE(s)', ', '.join(vuln['cves']) if vuln['cves'] and isinstance(vuln['cves'], list) else str(vuln['cves']) if vuln['cves'] else 'N/A'),
                ('Risk Factor', vuln['risk_factor']),
                ('Links/References', links_text)
            ]
            
            for row_idx, (label, value) in enumerate(info_data):
                row = info_table.rows[row_idx]
                row.cells[0].text = label
                row.cells[0].paragraphs[0].runs[0].font.bold = True
                row.cells[1].text = str(value)
            
            # Description section
            doc.add_heading('Description', level=4)
            desc_para = doc.add_paragraph(vuln['description'] or 'No description available.')
            
            # Plugin output section (if available)
            if vuln.get('plugin_output'):
                doc.add_heading('Technical Details', level=4)
                output_para = doc.add_paragraph()
                output_para.add_run('Plugin Output:').bold = True
                output_para.add_run('\n' + vuln['plugin_output'])
                output_para.style = 'Intense Quote'
            
            # Screenshots section (placeholder)
            doc.add_heading('Screenshots', level=4)
            screenshot_para = doc.add_paragraph()
            screenshot_para.add_run('[Screenshot placeholder - Insert relevant screenshots here]').italic = True
            screenshot_para2 = doc.add_paragraph()
            screenshot_para2.add_run('Note: Add screenshots showing the vulnerability exploitation or evidence here.').italic = True
            doc.add_paragraph()  # Add some space
            
            # Solution section
            doc.add_heading('Remediation', level=4)
            solution_para = doc.add_paragraph(vuln['solution'] or 'No solution provided.')
            
            # Affected hosts section
            doc.add_heading('Affected Hosts', level=4)
            
            if vuln['affected_hosts']:
                # Create table for affected hosts
                hosts_table = doc.add_table(rows=1, cols=5)
                hosts_table.style = 'Colorful List'
                hosts_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                
                # Header row
                header_cells = hosts_table.rows[0].cells
                header_cells[0].text = 'IP Address'
                header_cells[1].text = 'FQDN'
                header_cells[2].text = 'Operating System'
                header_cells[3].text = 'Port'
                header_cells[4].text = 'Protocol'
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add host rows
                for host in vuln['affected_hosts']:
                    row_cells = hosts_table.add_row().cells
                    row_cells[0].text = host['ip']
                    row_cells[1].text = host['fqdn'] or 'N/A'
                    row_cells[2].text = host['os'] or 'Unknown'
                    row_cells[3].text = host['port'] or 'N/A'
                    row_cells[4].text = host['protocol'] or 'N/A'
            else:
                doc.add_paragraph('No affected hosts identified.')
            
            # Add separator between vulnerabilities
            doc.add_paragraph()
            separator = doc.add_paragraph()
            separator.add_run('─' * 80).font.color.rgb = None  # Horizontal line
            doc.add_paragraph()
    
    def _create_default_template(self, doc: Document):
        """Create a default template if none exists"""
        
        # Title
        title = doc.add_heading('${ReportTitle}', 0)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run('This vulnerability assessment report was generated on ${ScanDate}. ')
        summary_para.add_run('The scan identified ${TotalVulnerabilities} vulnerabilities across ${TotalHosts} hosts.')
        
        # Summary table
        doc.add_heading('Vulnerability Summary', level=2)
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Colorful List'
        
        summary_data = [
            ('Critical', '${CriticalCount}'),
            ('High', '${HighCount}'),
            ('Medium', '${MediumCount}'),
            ('Low', '${LowCount}'),
            ('Informational', '${InfoCount}')
        ]
        
        # Header row
        summary_table.rows[0].cells[0].text = 'Severity'
        summary_table.rows[0].cells[1].text = 'Count'
        
        for i, (severity, count) in enumerate(summary_data):
            row = summary_table.rows[i + 1]
            row.cells[0].text = severity
            row.cells[1].text = count
        
        # Vulnerability details section
        doc.add_heading('Detailed Vulnerability Findings', level=1)
        doc.add_paragraph('The following section provides detailed information for each vulnerability identified during the assessment, organized by vulnerability type with affected hosts listed under each finding.')
        doc.add_paragraph('${VulnTable}')
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        doc.add_paragraph('Based on the vulnerability assessment findings, the following recommendations are provided:')
        
        rec_list = doc.add_paragraph()
        rec_list.style = 'List Bullet'
        rec_list.add_run('Prioritize remediation of Critical and High severity vulnerabilities')
        
        rec_list2 = doc.add_paragraph()
        rec_list2.style = 'List Bullet'
        rec_list2.add_run('Review and validate all findings with supporting screenshots')
        
        rec_list3 = doc.add_paragraph()
        rec_list3.style = 'List Bullet'
        rec_list3.add_run('Implement regular vulnerability scanning processes')
        
        rec_list4 = doc.add_paragraph()
        rec_list4.style = 'List Bullet'
        rec_list4.add_run('Establish patch management procedures for affected systems')
    
    def _format_date(self, date_string: str) -> str:
        """Format date string for display"""
        if not date_string:
            return datetime.now().strftime('%Y-%m-%d')
        
        try:
            # Try to parse common Nessus date formats
            for fmt in ['%a %b %d %H:%M:%S %Y', '%Y-%m-%d %H:%M:%S', '%Y-%m-%d']:
                try:
                    dt = datetime.strptime(date_string, fmt)
                    return dt.strftime('%Y-%m-%d')
                except ValueError:
                    continue
            
            # If parsing fails, return as-is or current date
            return date_string.split()[0] if ' ' in date_string else datetime.now().strftime('%Y-%m-%d')
            
        except Exception:
            return datetime.now().strftime('%Y-%m-%d')
    
    def add_logo(self, doc: Document, logo_path: str):
        """Add company logo to document"""
        if os.path.exists(logo_path):
            try:
                # Find a suitable place to insert logo (typically first paragraph)
                if doc.paragraphs:
                    paragraph = doc.paragraphs[0]
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(2))
                    logging.info(f"Added logo: {logo_path}")
            except Exception as e:
                logging.warning(f"Could not add logo {logo_path}: {e}")
        else:
            logging.warning(f"Logo file not found: {logo_path}") 